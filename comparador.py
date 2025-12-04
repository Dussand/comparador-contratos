import streamlit as st
import pandas as pd
import requests
from io import BytesIO, StringIO
from docx import Document
import time
import os


st.set_page_config(page_title="Comparador de Contratos", page_icon="📄")
st.title("📄 Comparador de Contratos – Kashio Legal")

# URL del webhook (usa la de producción si ya activaste el flujo)
N8N_WEBHOOK_URL_PRODUCTION = "https://operationskashio.app.n8n.cloud/webhook/compare_contracts"
#N8N_WEBHOOK_URL_TEST = 'http://localhost:5678/webhook-test/compare_contracts'
#N8N_WEBHOOK_URL_PRODUCTION_CORREO = "https://operationskashio.app.n8n.cloud/webhook/compare_contracts"


# ---------- FUNCIONES ----------
def leer_txt(file):
    """Lee archivos de texto plano (.txt)."""
    return file.getvalue().decode("utf-8")

def leer_docx(file):
    """Lee texto de un archivo .docx incluyendo párrafos y texto dentro de tablas."""
    doc = Document(BytesIO(file.getvalue()))
    texto = []

    # 1️⃣ Leer los párrafos normales
    for p in doc.paragraphs:
        if p.text.strip():
            texto.append(p.text.strip())

    # 2️⃣ Leer texto dentro de tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        texto.append(p.text.strip())

    # 3️⃣ Unir todo en un solo bloque de texto
    return "\n".join(texto)

def extraer_texto(file):
    """Detecta el tipo de archivo y obtiene el texto."""
    if file is None:
        return ""
    if file.name.endswith(".txt"):
        return leer_txt(file)
    elif file.name.endswith(".docx"):
        return leer_docx(file)
    else:
        return "⚠️ Formato no soportado (usa .docx o .txt)"
    
#----------GENERACION DE CORREO------------

def typing_effect(text, speed=0.003):
    placeholder = st.empty()
    output = ""
    for char in text:
        output += char
        placeholder.markdown(f"```\n{output}\n```")
        time.sleep(speed)

#===========================================================
# CODIGO PARA CONTRATOS ESTANDARIZADOS
#===========================================================

# # ---------- CARGA DE ARCHIVOS ----------

# col1, col2 = st.columns(2)
# with col1:
#     contrato_marco = st.file_uploader("**CONTRATO MARCO**", type=['docx', 'txt'])
# with col2: 
#     servicio_especifico = st.file_uploader('**SERVICIO ESPECIFICO**', type=['docx', 'txt'])

# # ---------- VISTA PREVIA ----------
# if contrato_marco:
#     text_contratoMarco = extraer_texto(contrato_marco)
#     st.subheader("📘 Contenido Contrato Marco:")
#     st.text_area("Contrato Marco", text_contratoMarco, height=250)

# else:
#     text_contratoMarco = ""

# if servicio_especifico:
#     text_servicioEspecifico = extraer_texto(servicio_especifico)
#     st.subheader("📕 Texto del Servicio Especifico:")
#     st.text_area("Servicio Especifico", text_servicioEspecifico, height=250)
# else:
#     text_servicioEspecifico = ""


# # ---------- ENVÍO A N8N ----------
# if st.button("🔍 Enviar a n8n para comparar"):
#     if not (contrato_marco and servicio_especifico):
#         st.warning("Por favor sube ambos contratos antes de continuar.")
#     else:
#         with st.spinner("Procesando en n8n..."):
#             try:
#                 # Enviar el texto procesado en formato JSON
#                 payload = {"contrato_marco": text_contratoMarco, "servicio_esp": text_servicioEspecifico}
#                 response = requests.post(N8N_WEBHOOK_URL_PRODUCTION, json=payload, timeout=600)

#                 if response.ok:
#                     data = response.json()
#                     #st.write("🧾 Respuesta completa del Webhook:", data)

#                     # Detectar formato flexible
#                     if isinstance(data, list):
#                         diferencias = data
#                     elif isinstance(data, dict):
#                         if isinstance(data.get("data"), list):
#                             diferencias = data["data"]
#                         elif isinstance(data.get("data"), dict):
#                             diferencias = [data["data"]]
#                         else:
#                             diferencias = []
#                     else:
#                         diferencias = []

#                     # Mostrar diferencias
#                     if diferencias:
#                         st.success(f"✅ Se detectaron {len(diferencias)} diferencia(s)")
#                         df = pd.DataFrame(diferencias)
#                         st.dataframe(df, use_container_width=True)
#                     else:
#                         st.info("No se detectaron diferencias relevantes o la respuesta vino vacía.")

#                 else:
#                     st.error(f"❌ Error {response.status_code}: {response.text}")

#             except requests.Timeout:
#                 st.error("⏱️ Tiempo de espera excedido. Intenta nuevamente.")
#             except Exception as e:
#                 st.error(f"❌ Error inesperado: {e}")

#===========================================================
# CODIGO PARA CONTRATOS DESACTUALIZADOS
#===========================================================


# ---------- CARGA DE ARCHIVOS ----------
col1, col2 = st.columns(2)
with col1:
    contrato_base = st.file_uploader("Contrato base", type=["docx", "txt"])
with col2:
    contrato_mod = st.file_uploader("Contrato modificado", type=["docx", "txt"])

# ---------- VISTA PREVIA ----------
if contrato_base:
    texto_base = extraer_texto(contrato_base)
    st.subheader("📘 Texto del contrato BASE:")
    st.text_area("Contenido base", texto_base, height=150)
else:
    texto_base = ""

if contrato_mod:
    texto_mod = extraer_texto(contrato_mod)
    st.subheader("📕 Texto del contrato MODIFICADO:")
    st.text_area("Contenido modificado", texto_mod, height=150)
else:
    texto_mod = ""

if st.button("Comparar contratos", use_container_width=True):

    if not (contrato_base and contrato_mod):
        st.warning("Por favor sube ambos contratos antes de continuar.")
    else:

        nombre_base = os.path.splitext(contrato_base.name)[0]
        nombre_mod = os.path.splitext(contrato_mod.name)[0]
        
        # ✅ BARRA DE PROGRESO
        progress_bar = st.progress(0)
        status_text = st.empty()

        try:
            status_text.text("🔍 Analizando diferencias...")
            progress_bar.progress(30)
            
            # Payload para el webhook
            payload = {"contrato_a": texto_base, "contrato_b": texto_mod, 'contrato_aName': nombre_base}

            status_text.text("🔍 Analizando diferencias...")
            progress_bar.progress(60)
            
            response = requests.post(N8N_WEBHOOK_URL_PRODUCTION, json=payload, timeout=600)

            status_text.text("✅ Procesando resultados...")
            progress_bar.progress(90)

            if response.ok:
                progress_bar.progress(100)
                time.sleep(0.3)  # Pausa breve para que se vea el 100%
                progress_bar.empty()
                status_text.empty()
                
                data = response.json()

                # ---------------------------------------------
                # INTERPRETACIÓN FLEXIBLE DE RESPUESTA
                # ---------------------------------------------
                # Detectar formato flexible según la nueva salida
                diferencias = []
                sin_diferencias = None

                # Caso A: lista → diferencias posibles
                if isinstance(data, list):
                    # Filtrar solo items que tengan tipo_cambio
                    diferencias = [d for d in data if d.get("tipo_cambio")]
                    if not diferencias:
                        # Si no hay tipo_cambio → se trata como sin diferencias
                        sin_diferencias = {"impacto_global": "Sin cambios"}

                # Caso B: un objeto → revisar impacto_global
                elif isinstance(data, dict):

                    # Caso "sin diferencias relevantes"
                    if data.get("tipo_cambio") == "Sin cambios":
                        sin_diferencias = data

                    # Caso {"data": [ ... ]}
                    elif isinstance(data.get("data"), list):
                        diferencias = [d for d in data["data"] if d.get("tipo_cambio")]
                        if not diferencias:
                            sin_diferencias = data

                    # Caso {"data": { ... }}
                    elif isinstance(data.get("data"), dict):
                        if data["data"].get("tipo_cambio"):
                            diferencias = [data["data"]]
                        else:
                            sin_diferencias = data["data"]

                # ---------------------------------------
                # MOSTRAR RESULTADOS
                # ---------------------------------------

                if sin_diferencias:
                    st.info("🟦 Sin diferencias relevantes en el contrato.")
                    #st.json(sin_diferencias)

                elif diferencias:
                    st.success(f"✅ Se detectaron {len(diferencias)} diferencias.")
                    st.dataframe(pd.DataFrame(diferencias), use_container_width=True)

                    # generar_correo = st.button("💬 Generar correo", use_container_width=True)


                    # if generar_correo:
                    #     payload_correo = {
                    #         "nombre_contrato_mod": nombre_mod,
                    #         "cliente": diferencias[0].get("cliente", "Cliente"),
                    #         "diferencias": diferencias,
                    #     }

                    #     response_correo = requests.post(N8N_WEBHOOK_URL_PRODUCTION_CORREO, json=payload_correo).json()

                    #     correo_generado = response_correo["correo"]

                    #     # ------------------------------
                    #     # EFECTO TYPING AQUÍ 🔥🔥
                    #     # ------------------------------
                    #     typing_effect(correo_generado, speed=0.003)

                else:
                    st.warning("⚠️ Respuesta no válida o inesperada.")


            else:
                progress_bar.empty()
                status_text.empty()
                st.error(f"❌ Error {response.status_code}: {response.text}")

        except requests.Timeout:
            progress_bar.empty()
            status_text.empty()
            st.error("⏱️ Tiempo de espera excedido. Intenta nuevamente.")
        except Exception as e:
            progress_bar.empty()
            status_text.empty()
            st.error(f"❌ Error inesperado: {e}")